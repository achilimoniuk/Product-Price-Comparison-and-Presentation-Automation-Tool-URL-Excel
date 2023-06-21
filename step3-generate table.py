# Script to create tables with statistics
import pandas as pd
import docx
from tqdm import tqdm

# SETUP: cuntries dictionary: 
# key- country,
# value- ending of the mismatch file (if the analysis wasn't done- leave blank)
countries = {'UK':''}

# Preparing output dataframes 
# Dataframe with numbers of lines
df = pd.DataFrame(columns=['Country', 
'Total # of Historical Sales Lines',
'Total # of checked Sales Lines',
'# of Lines with Successful Lookup',
'# of Lines with Failed Lookup',
'percentages'])

# Dataframe with numbers of lines based on the type of issue
df2 = pd.DataFrame(columns=['Country',
'Number of mismtached lines with List Price vs Contract ID issue',
'Number of mismtached lines with Different Contract IDs issue',
'Number of mismtached lines with Backdating issue',
'Number of mismtached lines with Other issue',
'percentages'])

# Dataframe with numbers of contracts
df3 = pd.DataFrame(columns=['Country', 
'Total # of Historical contracts',
'Total # of checked contracts',
'# of contracts with Successful Lookup',
'# of contracts with Failed Lookup',
'percentages'])

# Dataframe with numbers of contracts based on the type of issue
df4 = pd.DataFrame(columns=['Country',
'Number of mismtached contracts with List Price vs Contract ID issue',
'Number of mismtached contracts with Different Contract IDs issue',
'Number of mismtached contracts with Backdating issue',
'Number of mismtached contracts with Other issue',
'percentages'])

# Dataframe with numer of contracts based on the List Price type of issues (after analysis)
df5 = pd.DataFrame(columns=['Country', 
'values',
'percentages'])


for country in tqdm(countries):
    path = f'{country}/{country} Aug'
# Lines
    # Number of matched and mismatched lines
    pricesdf = pd.read_csv(f'{path}/files/all transactions {country}.csv')
    df.loc[len(df.index)] = [country,
    '', 
    pricesdf.shape[0],  
    pricesdf[pricesdf['Are prices matched?'] == 'yes'].shape[0], 
    pricesdf.loc[pricesdf['Are prices matched?'] == 'no'].shape[0],
    (pricesdf['Are prices matched?'].value_counts()/pricesdf['Are prices matched?'].count())*100]
    # Number of matched and mismatched lines based on type of issue
    mismatched = pd.read_excel(f'{path}/files/mismatched prices {country}{countries[country]}.xlsx')
    df2.loc[len(df2.index)] = [country,
    mismatched[mismatched['Issue type'] == 'List Price vs Contract ID'].shape[0],
    mismatched[mismatched['Issue type'] == 'Different Contract IDs'].shape[0],
    mismatched[mismatched['Issue type'] == 'Backdating issue'].shape[0],
    mismatched[mismatched['Issue type'] == 'Other issue'].shape[0],
    (mismatched['Issue type'].value_counts()/mismatched['Issue type'].count())*100]
# Contracts
    # Number of matched and mismatched contracts
    pricesdf = pricesdf[pricesdf['Contract ID from Excel'] != 'None']
    pricesdf['Contract ID from Excel'] = pd.to_numeric(pricesdf['Contract ID from Excel']) 
    pricesdf_matched = pricesdf[pricesdf['Are prices matched?'] == 'yes']
    pricesdf_mismatched = pricesdf[pricesdf['Are prices matched?'] == 'no']
    pricesdf_matched = pricesdf_matched.drop_duplicates(subset='Contract ID from Excel', keep="last")
    pricesdf_mismatched = pricesdf_mismatched.drop_duplicates(subset='Contract ID from Excel', keep="last")
    pricesdf = pricesdf.drop_duplicates(subset='Contract ID from Excel', keep="last")
    df3.loc[len(df3.index)] = [country,
    'contracts', 
    pricesdf.shape[0],  
    pricesdf_matched.shape[0], 
    pricesdf_mismatched.shape[0], 
    pricesdf_matched.shape[0]/(pricesdf_matched.shape[0]+pricesdf_mismatched.shape[0])*100]
    # Number of matched and mismatched contracts based on type of issue
    mismatched = mismatched[mismatched['Contract ID from Excel'] != 'None']
    mismatched['Contract ID from Excel'] = pd.to_numeric(mismatched['Contract ID from Excel']) 
    mismatched_listprice = mismatched[mismatched['Issue type'] == 'List Price vs Contract ID'] 
    mismatched_others = mismatched[mismatched['Issue type'] != 'List Price vs Contract ID'] 
    mismatched_listprice = mismatched_listprice.drop_duplicates(subset='Contract ID from Excel', keep="last")
    mismatched_others = mismatched_others.drop_duplicates(subset='Contract ID from Excel', keep="last")
    # number of contracts based on type of issue
    df4.loc[len(df4.index)] = [country,
    mismatched_listprice.shape[0],
    mismatched_others[mismatched_others['Issue type'] == 'Different Contract IDs'].shape[0],
    mismatched_others[mismatched_others['Issue type'] == 'Backdating issue'].shape[0],
    mismatched_others[mismatched_others['Issue type'] == 'Other issue'].shape[0],
    (mismatched['Issue type'].value_counts()/mismatched['Issue type'].count())*100]


# opening new document
doc = docx.Document()

# function for creating tables from dataframes
def save_doc(df, text):
    doc.add_paragraph('')
    doc.add_paragraph(f'{text}') # adding text in front of the table
    t = doc.add_table(df.shape[0]+1, df.shape[1], style="Table Grid") # adding table with extra row for headers
    # adding header row
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    # adding the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

# using function for adding tables for customer, program and lineitem
save_doc(df, 'number of all lines')
save_doc(df2, 'number of all lines- type of issue')
save_doc(df3, 'number of all contracts')
save_doc(df4, 'number of all contracts- type of issue')


# saving the document
doc.save('output_file.docx')